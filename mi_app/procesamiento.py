import os
import re
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import cv2
import numpy as np
from openai import OpenAI
from django.conf import settings
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from datetime import date
import fitz 

client = OpenAI(
    api_key=settings.OPENAI_API_KEY,
    base_url="https://api.together.xyz/v1"
)

# ✅ Utilidades
def limpiar_texto(texto):
    texto = texto.replace("\n", " ")
    texto = re.sub(r'[^A-Za-z0-9ÁÉÍÓÚáéíóúüÜñÑ¿?!,.;:\-/\s]', '', texto)
    texto = re.sub(r'\s+', ' ', texto).strip()
    return texto

def preprocesar_imagen(image_pil):
    img = np.array(image_pil.convert('RGB'))
    img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 15, 10)
    kernel = np.ones((1, 1), np.uint8)
    img = cv2.morphologyEx(img, cv2.MORPH_OPEN, kernel)
    img = cv2.morphologyEx(img, cv2.MORPH_CLOSE, kernel)
    return Image.fromarray(img)

# ✅ IA: Limpieza texto OCR
def limpiar_texto_con_llama3(texto_raw):
    prompt = (
        "Recibiste un texto extraído de un PDF mediante OCR. "
        "Corrige faltas de ortografía y puntuación, elimina basura, "
        "NO resumas ni cambies palabras como 'y/o'. "
        "Si 'yo' aparece mal, cámbialo por 'o bien'. "
        "Devuelve solo texto limpio sin encabezados.\n\nTexto OCR original:\n\n"
        + texto_raw
    )
    try:
        response = client.chat.completions.create(
            model="meta-llama/Llama-3-70b-chat-hf",
            messages=[
                {"role": "system", "content": "Eres un corrector de texto oficial en español."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[❌] Error IA: {e}")
        return texto_raw

# ✅ IA: Corrección preguntas OCR
def corregir_preguntas_con_llama3(lista_preguntas):
    bloque = "\n".join(lista_preguntas)
    prompt = (
        "Corrige estas preguntas OCR en español sin parafrasear. "
        "Corrige ortografía y puntuación, mantén 'y/o'. Si aparece 'yo' "
        "como error, reemplaza por 'o bien'. Devuelve solo la lista sin encabezados.\n\n"
        "Preguntas:\n\n"
        + bloque
    )
    try:
        response = client.chat.completions.create(
            model="meta-llama/Llama-3-70b-chat-hf",
            messages=[
                {"role": "system", "content": "Eres un corrector de estilo en español."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip().splitlines()
    except Exception as e:
        print(f"[❌] Error IA preguntas: {e}")
        return lista_preguntas

# ✅ 1) OCR PDF escaneado y limpieza
def ocr_pdf_escaneado(pdf_path, output_txt_path):
    paginas = convert_from_path(pdf_path, dpi=300)
    texto_total = ""
    numero_solicitud = "[NO ENCONTRADO]"

    for i, pagina in enumerate(paginas):
        imagen_preprocesada = preprocesar_imagen(pagina)
        texto = pytesseract.image_to_string(imagen_preprocesada, lang='spa')

        if numero_solicitud == "[NO ENCONTRADO]":
            match_solicitud = re.search(r"\b(33\d+)\b", texto)
            if match_solicitud:
                numero_solicitud = match_solicitud.group(1).strip()

        texto_limpio = limpiar_texto(texto)
        texto_total += texto_limpio + "\n\n"

    texto_corregido = limpiar_texto_con_llama3(texto_total)

    if numero_solicitud != "[NO ENCONTRADO]":
        texto_corregido += f"\n\nNúmero de solicitud detectado: {numero_solicitud}"

    with open(output_txt_path, 'w', encoding='utf-8') as f:
        f.write(texto_corregido)

    print(f"[✔] OCR + IA guardado en: {output_txt_path}")

# ✅ 2) Extracción información clave
def extraer_informacion_especifica(ruta_txt, ruta_txt_extraido):
    with open(ruta_txt, 'r', encoding='utf-8') as f:
        contenido = f.read()

    contenido = re.sub(
        r'(?m)^[ \t]*[aAiI]\)\s*',
        '1. ',
        contenido
    )

    contenido = re.sub(
        r'(?m)^[ \t]*(\d+)\)\s*',
        r'\1. ',
        contenido
    )

    match_fecha = re.search(
        r"(Ciudad de México,\s*a\s*\d{1,2}\s+de\s+[a-zA-ZñÑ]+\s+de\s+\d{1,5})",
        contenido, re.IGNORECASE)
    texto_fecha = match_fecha.group(1).strip() if match_fecha else "[FECHA NO ENCONTRADA]"

    match_oficio = re.search(r"(Oficio Núm\.?\s*[\w\-/]+|No\.\s*\d+)", contenido, re.IGNORECASE)
    if match_oficio:
        texto_oficio = match_oficio.group(1).strip() if match_oficio else "[OFICIO NO ENCONTRADO]"
        texto_oficio = re.sub(
            r'([A-Z]+-[A-Z]+)(\d{2})(\d{3})(\d{2})',
            r'\1/\2/\3/\4',
            texto_oficio
        ).replace('//', '/')
    else:
        texto_oficio = "[OFICIO NO ENCONTRADO]"
    
    match_solicitud = re.search(r"\b(33\d+)\b", contenido)
    texto_solicitud = f"No. {match_solicitud.group(1).strip()}" if match_solicitud else "[SOLICITUD NO ENCONTRADA]"

    patron_preguntas = re.compile(
        r"^[ \t]*(\d+\.\s+.*?)(?=^[ \t]*\d+\.\s|\n\n|$)", 
        re.DOTALL | re.MULTILINE
    )
    
    lista_raw = []
    
    for m in patron_preguntas.finditer(contenido):
        pregunta = m.group(1).strip()
        pregunta = re.sub(r'\s+', ' ', pregunta)
        pregunta = re.sub(r'\byo\b', 'o bien', pregunta, flags=re.IGNORECASE)
        lista_raw.append(pregunta)

    if lista_raw:
        processed_questions = corregir_preguntas_con_llama3(lista_raw)
        numeros = []
        for p in processed_questions:
            match = re.match(r'^(\d+)\.', p)
            if match:
                numeros.append(int(match.group(1)))
        
        if numeros and numeros != list(range(1, len(numeros)+1)):
            print(f"Advertencia: Números de preguntas podrían estar desordenados: {numeros}")
    else:
        processed_questions = []

    texto_preguntas = "\n".join(processed_questions) if processed_questions else "[PREGUNTAS NO ENCONTRADAS]"

    texto_extraido = "\n\n".join([
        f"Fecha: {texto_fecha}",
        f"Número de oficio: {texto_oficio}",
        f"Número de solicitud: {texto_solicitud}",
        "Preguntas:",
        texto_preguntas
    ])

    with open(ruta_txt_extraido, 'w', encoding='utf-8') as f:
        f.write(texto_extraido)

    print(f"[✔] Extracción + corrección completada. Guardado en: {ruta_txt_extraido}")

# ✅ 3) Construccion Word Oficios
def exportar_a_word(ruta_txt: str, ruta_docx: str, nombre: str, cargo: str, cuerpo_respuesta: str, ccp_adicional: str) -> None:
    with open(ruta_txt, encoding="utf-8") as f:
        raw = f.read().strip()
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{2,}", "\n", raw)

    m_fecha_simple = re.search(
        r"a\s+(\d{1,2}\s+de\s+[a-zA-Z]+(?:\s+de\s+\d{4})?)",
        raw,
        re.IGNORECASE
    )
    fecha_simple = m_fecha_simple.group(1) if m_fecha_simple else "[fecha no encontrada]"

    m_num_solic = re.search(r"\b(33\d+)\b", raw)
    num_solic = m_num_solic.group(1).strip() if m_num_solic else "[NÚMERO NO ENCONTRADO]"


    m_oficio = re.search(r"(Oficio Núm\.?\s*[\w\-/]+|No\.\s*\d+)", raw, re.IGNORECASE)
    oficio_texto = m_oficio.group().strip() if m_oficio else "[OFICIO NO ENCONTRADO]"

    cuerpo = (
        f"Hago referencia al {oficio_texto} "
        f"de fecha {fecha_simple}, mediante el cual se informa de la recepción "
        f"de la solicitud de información pública No. {num_solic}, "
        "a través de la Plataforma Nacional de Transparencia, y en donde "
        "el particular solicita lo siguiente:"
    )

    patron_preguntas = r"Preguntas:\s*(.*)"
    m_pregs = re.search(patron_preguntas, raw, re.DOTALL | re.IGNORECASE)

    preguntas_raw = []
    for linea in raw.splitlines():
        linea = linea.strip()
        m = re.match(r"^(\d+)\.\s+(.*\S)", linea)
        if m:
            numero = int(m.group(1))
            texto = m.group(2).strip()
            preguntas_raw.append((numero, texto))

    expected = list(range(1, len(preguntas_raw) + 1))
    actual = [num for num, _ in preguntas_raw]
    if actual != expected:
        raise ValueError(f"Números de preguntas desordenados: {actual}")


    doc = Document()

    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    hoy = date.today()
    fecha_hoy = (
        f"Ciudad de México, a {hoy.day} de "
        f"{meses[hoy.month - 1]} de {hoy.year}"
    )

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans"
    normal.font.size = Pt(9)
    rfonts = normal.element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Noto Sans")
    rfonts.set(qn("w:hAnsi"), "Noto Sans")

    def _agrega_parrafo(texto: str):
        p = doc.add_paragraph(texto.strip())
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.line_spacing = 1.15
        return p

    doc.add_paragraph()
    p = doc.add_paragraph(fecha_hoy)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.runs[0]
    run.font.name = "Noto Sans"
    run.font.size = Pt(9)
    r = run._element.rPr.rFonts
    r.set(qn("w:ascii"), "Noto Sans")
    r.set(qn("w:hAnsi"), "Noto Sans")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph("Oficio Núm. INFOTEC-DAIC-GI-SIG-XXXXX")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.runs[0]
    run.font.name = "Noto Sans SemiBold"
    run.font.size = Pt(9)
    run.bold = True
    r = run._element.rPr.rFonts
    r.set(qn("w:ascii"), "Noto Sans SemiBold")
    r.set(qn("w:hAnsi"), "Noto Sans SemiBold")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    for linea in [f"{nombre}", f"{cargo}", "PRESENTE"]:
        p = doc.add_paragraph(linea)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.runs[0]
        run.font.name = "Noto Sans ExtraBold"
        run.font.size = Pt(9)
        run.bold = True
        r = run._element.rPr.rFonts
        r.set(qn("w:ascii"), "Noto Sans ExtraBold")
        r.set(qn("w:hAnsi"), "Noto Sans ExtraBold")
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()

    _agrega_parrafo(cuerpo)
    doc.add_paragraph()

    total_q = len(preguntas_raw)

    for idx, (_num, txt) in enumerate(preguntas_raw, start=1):
        p_preg = doc.add_paragraph()
        p_preg.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if idx == 1:
            run_open = p_preg.add_run('“')
            run_open.bold = True
            run_open.font.name = "Noto Sans"
            run_open.font.size = Pt(9)

        run_preg = p_preg.add_run(f"{_num}. {txt}")
        run_preg.bold = True
        run_preg.font.name = "Noto Sans"
        run_preg.font.size = Pt(9)

        if idx == total_q:
            run_close = p_preg.add_run('”')
            run_close.bold = True
            run_close.font.name = "Noto Sans"
            run_close.font.size = Pt(9)

            run_sic = p_preg.add_run(' (Sic).')
            run_sic.font.name = "Noto Sans"
            run_sic.font.size = Pt(9)
    doc.add_paragraph()

    cierre_mod = cuerpo_respuesta.strip() or (
        "Sobre el particular, después de realizada una búsqueda exhaustiva y razonable "
        "en los archivos físicos y electrónicos que obran en la Dirección Adjunta de Innovación y Conocimiento, "
        "así como de la consulta al personal adscrito a la misma, se comunica lo siguiente:"
    )
    _agrega_parrafo(cierre_mod)
    doc.add_paragraph()

    for idx, (_num, txt) in enumerate(preguntas_raw, start=1):
        if idx == total_q:
            prefix = "Finalmente, con relación a “[…] "
        elif idx % 2 == 1:
            prefix = "Con relación a “[…] "
        else:
            prefix = "Con respecto a “[…] "
        suffix = " […]” (Sic)," 

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15

        run_pre = p.add_run(prefix)
        run_pre.font.name = "Noto Sans"
        run_pre.font.size = Pt(9)

        run_bold = p.add_run(f"{_num}. {txt}")
        run_bold.bold = True
        run_bold.font.name = "Noto Sans"
        run_bold.font.size = Pt(9)

        run_suf = p.add_run(suffix)
        run_suf.font.name = "Noto Sans"
        run_suf.font.size = Pt(9)
        doc.add_paragraph()

    doc.add_paragraph()
    _agrega_parrafo("Sin otro particular, aprovecho la ocasión para enviarle un cordial saludo.")

    doc.add_paragraph() 
    firma_text = ("ATENTAMENTE\n\n\n\n\nDR. JUAN ANTONIO VEGA GARFIAS\n"
                  "SUBGERENTE DE INNOVACIÓN GUBERNAMENTAL\n"
                  "ENLACE EN LA DAIC PARA LA ATENCIÓN DE SOLICITUDES DE INFORMACIÓN PÚBLICA")
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_firma = p_firma.add_run(firma_text)
    run_firma.bold = True
    run_firma.font.name = "Noto Sans ExtraBold"
    run_firma.font.size = Pt(9)
    for _ in range(5):
        doc.add_paragraph()

    lineas_ccp = [
        "SIN ANEXOS.",
        "C.c.p.-  Mtro. Felipe Alfonso Delgado Castillo, Encargado de la Dirección Adjunta de Innovación y Conocimiento. Para su conocimiento.",
        "C.c.p.-  Mtra. Analy Mendoza Rosales, Encargada de la Gerencia de Capital Humano. Para su conocimiento.",  
    ]

    if ccp_adicional.strip():
        lineas_ccp.append(f"C.c.p.-  {ccp_adicional.strip()}. Para su conocimiento.")

    total_lineas = len(lineas_ccp)

    for idx, linea in enumerate(lineas_ccp):
        p = doc.add_paragraph()
        if idx == total_lineas - 1 and ccp_adicional.strip():
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(0)  

        if linea.strip().upper().startswith("SIN ANEXOS"):
            run = p.add_run(linea.strip())
            run.font.name = "Noto Sans"
            run.font.size = Pt(7.5)
        else:
            m = re.match(r"(C\.c\.p\.-\s+)(.*?)(,.*)", linea.strip())
            if m:
                ccp, nombre, resto = m.groups()
                run_ccp = p.add_run(ccp)
                run_ccp.bold = True
                run_ccp.font.name = "Noto Sans"
                run_ccp.font.size = Pt(7)
                run_nombre = p.add_run(nombre)
                run_nombre.bold = True
                run_nombre.font.name = "Noto Sans"
                run_nombre.font.size = Pt(7)
                run_resto = p.add_run(resto)
                run_resto.font.name = "Noto Sans"
                run_resto.font.size = Pt(7)
            else:
                run = p.add_run(linea.strip())
                run.font.name = "Noto Sans"
                run.font.size = Pt(7)

    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.left_indent = Cm(-2.69)
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    run = paragraph.add_run()
    p.paragraph_format.left_indent = Cm(-2.69)  
    run.add_picture("img/principio.png", width=Cm(19.98), height=Cm(2.39))    
    for linea in [
        "Dirección Adjunta de Innovación y Conocimiento",
        "Gerencia de Innovación",
        "Subgerencia de Innovación Gubernamental"]:
        p = header.add_paragraph(linea)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = p.runs[0]
        run.font.name = 'Noto Sans'
        run.font.size = Pt(9)
        run.bold = True
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:ascii'), 'Noto Sans')
        rFonts.set(qn('w:hAnsi'), 'Noto Sans')
        p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.space_after = Pt(0)

    footer = section.footer

    for p in footer.paragraphs:
        p.clear()

    p = footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.left_indent = Cm(-2.80)
    p.paragraph_format.space_before = Pt(20)  
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run()
    run.add_picture("img/final.png", width=Cm(20.48), height=Cm(2.1))

    doc.save(ruta_docx)
    print(f"✅ Documento Word creado: {ruta_docx}")

# ✅ 4) Construccion Word Simple
def _aplicar_estilo_parrafo(parrafo):
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in parrafo.runs:
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)

def convertir_pdf_escaneado_a_word(pdf_path: str, output_path: str) -> str:
    print("Iniciando conversión de PDF escaneado...")
    paginas_img = convert_from_path(pdf_path, dpi=300)
    texto_completo_raw = ""

    for i, pagina_img in enumerate(paginas_img, start=1):
        print(f"Procesando página {i}/{len(paginas_img)} con OCR...")
        img_proc = preprocesar_imagen(pagina_img)
        texto_pagina = pytesseract.image_to_string(img_proc, lang='spa')
        texto_completo_raw += texto_pagina + "\n\n" 
    print("Limpiando texto extraído con IA...")
    texto_limpio_final = limpiar_texto_con_llama3(texto_completo_raw)

    doc = Document()
    for linea in texto_limpio_final.split('\n'):
        if linea.strip(): 
            p = doc.add_paragraph(linea)
            _aplicar_estilo_parrafo(p)

    print(f"Guardando documento en: {output_path}")
    doc.save(output_path)
    return output_path

def convertir_pdf_digital_a_word(pdf_path: str, output_path: str) -> str:
    print("Iniciando conversión de PDF digital...")
    pdf = fitz.open(pdf_path)
    texto_completo = ""
    for i, pagina in enumerate(pdf, start=1):
        print(f"Extrayendo texto de la página {i}/{len(pdf)}...")
        texto_completo += pagina.get_text("text") + "\n\n"

    pdf.close()
    doc = Document()
    for linea in texto_completo.split('\n'):
        if linea.strip():
            p = doc.add_paragraph(linea)
            _aplicar_estilo_parrafo(p)
            
    print(f"Guardando documento en: {output_path}")
    doc.save(output_path)
    return output_path

def convertir_pdf_imagenes_a_word(pdf_path: str, output_path: str) -> str:
    print("Iniciando conversión de PDF con imágenes...")
    doc = Document()
    pdf = fitz.open(pdf_path)

    for i, pagina in enumerate(pdf, start=1):
        print(f"Procesando página {i}/{len(pdf)}...")
        
        texto_pagina = pagina.get_text("text")
        if texto_pagina.strip():
            for linea in texto_pagina.split('\n'):
                if linea.strip():
                    p = doc.add_paragraph(linea)
                    _aplicar_estilo_parrafo(p)

        imagenes = pagina.get_images(full=True)
        if imagenes:
            print(f"Encontradas {len(imagenes)} imágenes en la página {i}.")
            for img_index, img in enumerate(imagenes, start=1):
                xref = img[0]
                try:
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    img_temp_path = os.path.join("media", f"temp_page{i}_img{img_index}.{base_image['ext']}")
                    with open(img_temp_path, "wb") as f:
                        f.write(image_bytes)
                    
                    doc.add_picture(img_temp_path, width=Cm(15)) 
                    os.remove(img_temp_path) 

                except Exception as e:
                    print(f" [!] Error al procesar imagen en página {i}: {e}")

    print(f"Guardando documento en: {output_path}")
    doc.save(output_path)
    pdf.close()
    return output_path